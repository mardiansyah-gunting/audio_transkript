"""
Transcribe long audio files using speech recognition, with chunking for length limits.
Requires ffmpeg on PATH for pydub to decode many formats.
"""

import argparse
import os
import sys

import speech_recognition as sr
from docx import Document
from pydub import AudioSegment
from pydub.utils import make_chunks


def transcribe_long_audio(path_audio: str, language: str = "id-ID") -> str:
    """Split audio into chunks, transcribe each, return concatenated text."""
    recognizer = sr.Recognizer()
    full_text: list[str] = []

    print(f"Loading file: {path_audio}")
    try:
        if path_audio.lower().endswith(".m4a"):
            audio = AudioSegment.from_file(path_audio, format="m4a")
        else:
            audio = AudioSegment.from_file(path_audio)

        audio = audio.set_channels(1).set_frame_rate(16000)

    except Exception as e:
        return f"Error loading or converting audio: {e}"

    chunk_length_ms = 45000
    chunks = make_chunks(audio, chunk_length_ms)

    print(f"Audio split into {len(chunks)} segments for processing.\n")

    for i, chunk in enumerate(chunks):
        temp_wav = f"temp_chunk_{i}.wav"
        chunk.export(temp_wav, format="wav")

        print(f"Processing segment {i + 1} of {len(chunks)}...")

        try:
            with sr.AudioFile(temp_wav) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data, language=language)
                full_text.append(text)

        except sr.UnknownValueError:
            print(f" -> Segment {i + 1}: No clear speech detected.")
        except sr.RequestError as e:
            print(f" -> Segment {i + 1}: Recognition service error ({e})")
        except Exception as e:
            print(f" -> Segment {i + 1}: Unexpected error ({e})")

        finally:
            if os.path.exists(temp_wav):
                os.remove(temp_wav)

    result = " ".join(full_text)

    if not result:
        return ""

    return result


def save_to_word(text: str, path_out: str) -> None:
    doc = Document()
    doc.add_heading("Audio transcription", 0)
    doc.add_paragraph(text)
    doc.save(path_out)
    print(f"\nDone. Word file saved to: {path_out}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Transcribe a long audio file and save the result as a .docx file."
    )
    parser.add_argument(
        "input_audio",
        help="Path to the input audio file (e.g. .m4a, .mp3, .wav)",
    )
    parser.add_argument(
        "output_docx",
        help="Path for the output Word document (.docx)",
    )
    parser.add_argument(
        "--language",
        default="id-ID",
        help="BCP-47 language tag for recognition (default: id-ID for Indonesian)",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()

    if not os.path.isfile(args.input_audio):
        print(f"Input file not found: {args.input_audio}", file=sys.stderr)
        return 1

    out_dir = os.path.dirname(os.path.abspath(args.output_docx))
    if out_dir and not os.path.isdir(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    print("=" * 50)
    print("Starting transcription...")
    result = transcribe_long_audio(args.input_audio, language=args.language)

    if result and not result.startswith("Error "):
        save_to_word(result, args.output_docx)
    elif result.startswith("Error "):
        print(result, file=sys.stderr)
        return 1
    else:
        print("\nStopped: no text could be extracted from the audio.")
        return 1

    print("=" * 50)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
